#!/usr/bin/env python3
"""
Fix line-break patterns of a DOCX transcript (doc1) so that its lines follow
the line-break pattern of a reference DOCX (doc2, e.g. an SRT export).

How it works
------------
1. Both documents are read paragraph-by-paragraph and flattened into word
   streams (musical-note markers like '♪' are tracked separately per line).
2. The two word streams are aligned:
   - difflib.SequenceMatcher on normalised words gives exact "anchor" blocks
     (≈95% of words match after normalisation),
   - the small unmatched gap regions in between are aligned with a
     Needleman-Wunsch DP using per-word fuzzy similarity.
3. Every line boundary of doc2 is mapped through the alignment onto a cut
   position in doc1's word stream.  Doc1's words are then re-grouped using
   those cuts, so doc1 keeps ITS OWN words/spelling but gets doc2's breaks.
4. Reference lines that have no counterpart in doc1 (extra lines in doc2)
   simply map to an empty span and are left untouched / skipped.
5. The result is written as a new DOCX (one paragraph per line, blank
   paragraph between lines, same font as the input).

Usage:
    python fix_linebreaks.py wrong.docx reference.docx fixed_output.docx
"""

import re
import sys
import difflib

import docx
from docx.shared import Pt


# ----------------------------------------------------------------------
# helpers
# ----------------------------------------------------------------------

NOTE = "♪"


def normalise(word: str) -> str:
    """Normalise a word for alignment purposes (punctuation/diacritic noise)."""
    w = re.sub(r"[♪\[\]\(\)।॥,\.!\?\"'…:;\-–—_0-9]+", "", word)
    w = w.replace("ँ", "ं")          # chandrabindu ~ anusvara
    w = w.replace("\u200d", "").replace("\u200c", "")  # ZWJ / ZWNJ
    return w


def read_lines(path_or_file):
    """Return list of non-empty paragraph texts from a docx."""
    d = docx.Document(path_or_file)
    return [p.text.strip() for p in d.paragraphs if p.text.strip()]


def line_to_tokens(line: str):
    """Split a line into tokens, removing standalone note marks.

    Returns (tokens, has_note) where has_note is True when the line is a
    lyric line wrapped in ♪ ... ♪.
    """
    raw = line.split()
    has_note = any(t.strip(NOTE) == "" and NOTE in t for t in raw) or \
        line.startswith(NOTE) or line.endswith(NOTE)
    tokens = []
    for t in raw:
        t2 = t.strip(NOTE).strip()
        if t2:
            tokens.append(t2)
    return tokens, has_note


def gap_align(a, b):
    """Needleman-Wunsch alignment of two short normalised word lists.

    Returns list of ops: ('m', i, j) aligned pair, ('d', i, None) word only
    in a, ('i', None, j) word only in b.
    """
    n, m = len(a), len(b)
    GAP = 0.0
    # score matrix
    S = [[0.0] * (m + 1) for _ in range(n + 1)]
    P = [[None] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1):
        S[i][0] = S[i - 1][0] + GAP
        P[i][0] = "d"
    for j in range(1, m + 1):
        S[0][j] = S[0][j - 1] + GAP
        P[0][j] = "i"
    for i in range(1, n + 1):
        ai = a[i - 1]
        for j in range(1, m + 1):
            bj = b[j - 1]
            sim = difflib.SequenceMatcher(None, ai, bj).ratio() - 0.45
            best = S[i - 1][j - 1] + sim
            ptr = "m"
            if S[i - 1][j] + GAP > best:
                best, ptr = S[i - 1][j] + GAP, "d"
            if S[i][j - 1] + GAP > best:
                best, ptr = S[i][j - 1] + GAP, "i"
            S[i][j], P[i][j] = best, ptr
    # traceback
    ops = []
    i, j = n, m
    while i > 0 or j > 0:
        p = P[i][j]
        if p == "m":
            ops.append(("m", i - 1, j - 1))
            i, j = i - 1, j - 1
        elif p == "d":
            ops.append(("d", i - 1, None))
            i -= 1
        else:
            ops.append(("i", None, j - 1))
            j -= 1
    ops.reverse()
    return ops


def build_cut_map(w1n, w2n):
    """For each boundary position j (0..len(w2n)) in the reference stream,
    return c[j] = position in stream 1 such that w1[:c[j]] corresponds to
    w2[:j]."""
    n1, n2 = len(w1n), len(w2n)
    c = [0] * (n2 + 1)

    sm = difflib.SequenceMatcher(None, w1n, w2n, autojunk=False)
    blocks = sm.get_matching_blocks()  # ends with dummy (n1, n2, 0)

    pi, pj = 0, 0  # current position in each stream
    for blk in blocks:
        bi, bj, size = blk.a, blk.b, blk.size
        # ---- gap region [pi:bi] x [pj:bj] ----
        if bj > pj or bi > pi:
            ops = gap_align(w1n[pi:bi], w2n[pj:bj])
            ci = pi
            for op, oi, oj in ops:
                if op == "m":
                    c[pj + oj] = ci
                    ci = pi + oi + 1
                    c[pj + oj + 1] = ci
                elif op == "d":
                    ci = pi + oi + 1
                elif op == "i":
                    c[pj + oj + 1] = ci
        # ---- matching block ----
        for k in range(size):
            c[bj + k] = bi + k
            c[bj + k + 1] = bi + k + 1
        pi, pj = bi + size, bj + size

    c[n2] = n1  # consume everything that remains at the very end
    # enforce monotonicity
    for j in range(1, n2 + 1):
        if c[j] < c[j - 1]:
            c[j] = c[j - 1]
    return c


# ----------------------------------------------------------------------
# main fixing routine
# ----------------------------------------------------------------------

def fix_linebreaks(doc1_path, doc2_path):
    """Return list of (text, has_note) output lines for the fixed doc1."""
    lines1 = read_lines(doc1_path)
    lines2 = read_lines(doc2_path)

    # flatten doc1, remembering for EACH word whether it was inside a
    # ♪ ... ♪ lyric line in doc1 itself (we keep doc1's music notes,
    # never doc2's)
    w1 = []
    noted1 = []
    for ln in lines1:
        toks, has_note = line_to_tokens(ln)
        w1.extend(toks)
        noted1.extend([has_note] * len(toks))

    # flatten doc2, remembering each line's token span + note flag
    w2 = []
    spans2 = []  # (start, end, has_note)
    for ln in lines2:
        toks, has_note = line_to_tokens(ln)
        spans2.append((len(w2), len(w2) + len(toks), has_note))
        w2.extend(toks)

    w1n = [normalise(t) for t in w1]
    w2n = [normalise(t) for t in w2]

    c = build_cut_map(w1n, w2n)

    out = []
    prev_end = 0
    for (a, b, ref_note) in spans2:
        # start exactly where the previous line ended so that no word of
        # doc1 is ever lost; end at the mapped cut position
        s = prev_end
        e = max(c[b], s)
        toks = w1[s:e]
        prev_end = e
        if not toks:
            # extra reference line with no counterpart in doc1 -> skip
            continue
        # a line is a lyric line ONLY if these words were inside ♪ ... ♪
        # in doc1 itself (doc2's note marks are ignored completely)
        n_noted = sum(1 for f in noted1[s:e] if f)
        has_note = n_noted * 2 >= len(toks) and n_noted > 0
        text = " ".join(toks)
        out.append((text, has_note))
    # any tail words of doc1 not consumed (shouldn't happen, but be safe)
    if prev_end < len(w1):
        s = prev_end
        n_noted = sum(1 for f in noted1[s:] if f)
        has_note = n_noted * 2 >= len(w1) - s and n_noted > 0
        out.append((" ".join(w1[s:]), has_note))
    return out


def write_docx(lines, out_path, font_name="Courier New", font_size=None):
    d = docx.Document()
    style = d.styles["Normal"]
    style.font.name = font_name
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    if font_size:
        style.font.size = Pt(font_size)
    first = True
    for text, has_note in lines:
        if not first:
            d.add_paragraph("")  # exactly ONE empty line between two lines
        first = False
        # ♪ marks come from doc1's own lyric lines only (never from doc2)
        if has_note:
            text = f"{NOTE}  {text}  {NOTE}"
        d.add_paragraph(text)
    d.save(out_path)


def main():
    if len(sys.argv) != 4:
        print(__doc__)
        sys.exit(1)
    doc1, doc2, out = sys.argv[1], sys.argv[2], sys.argv[3]
    lines = fix_linebreaks(doc1, doc2)
    write_docx(lines, out)
    print(f"Wrote {len(lines)} lines -> {out}")


if __name__ == "__main__":
    main()